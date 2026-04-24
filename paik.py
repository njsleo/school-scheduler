import streamlit as st
import pandas as pd
import random
import copy
import math
import io
import re
import hashlib
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ================= 页面全局设置 =================
st.set_page_config(page_title="英华学校教务综合平台", layout="wide", page_icon="🏫")

# 初始化系统缓存
if 'schedule_result' not in st.session_state: st.session_state['schedule_result'] = None
if 'schedule_classes' not in st.session_state: st.session_state['schedule_classes'] = None
if 'exam_result' not in st.session_state: st.session_state['exam_result'] = None
if 'exam_slips' not in st.session_state: st.session_state['exam_slips'] = None
if 'slip_export_data' not in st.session_state: st.session_state['slip_export_data'] = None

# ================= 侧边栏导航 =================
st.sidebar.title("🛠️ 英华教务工作台")
app_mode = st.sidebar.radio("请选择功能模块：", ["📅 智能排课系统", "📝 全科考场编排"])
st.sidebar.markdown("---")

# =====================================================================
#                          模块一：智能排课系统 (全量未删减)
# =====================================================================
if app_mode == "📅 智能排课系统":
    DAYS = ['周一', '周二', '周三', '周四', '周五', '周六', '周日']
    SLOTS = [
        "早自习\n7:00-7:40", "第一节\n8:10-8:55", "第二节\n9:05-9:50", "第三节\n10:15-11:00", 
        "第四节\n11:00-11:45", "第五节\n14:00-14:45", "第六节\n14:55-15:40", "第七节\n15:50-16:35", 
        "第八节\n16:45-18:15", "第九节\n19:00-20:30", "第十节\n20:50-22:30"
    ]
    SUBJECT_COLORS = {
        "走班": "#FFFF99", "语文": "#FFE4E1", "数学": "#E4F1FF", "英语": "#E0FFFF", 
        "物理": "#E8F5E9", "化学": "#FFFACD", "生物": "#FCE4EC", "政治": "#FFDAB9", 
        "历史": "#FFE4B5", "地理": "#E1BEE7", "体育": "#C8E6C9", "音乐": "#E1BEE7", 
        "美术": "#FFCCBC", "班会": "#B3E5FC", "晚自习": "#F8F9FA", "自习": "#F8F9FA",
        "信息": "#E0F7FA", "通用": "#F5F5F5"
    }

    with st.sidebar:
        st.header("⚙️ 排课参数设置")
        class_prefix = st.text_input("班级前缀 (如: 高一 / 高三理科)", value="高三理科")
        class_suffix = st.text_input("班级后缀", value="班")
        include_sunday_evening = st.checkbox("✅ 开启周日晚修 (高三模式)", value=True)

    class CourseScheduler:
        def __init__(self, teacher_df, prefix, suffix, include_sunday):
            self.teacher_data = teacher_df.fillna("") 
            self.classes = set()
            self.teacher_busy = {} 
            self.prefix = prefix
            self.suffix = suffix
            self.evening_days = ['周一', '周二', '周三', '周四', '周五', '周日'] if include_sunday else ['周一', '周二', '周三', '周四', '周五']

        def parse_tasks(self):
            tasks, zouban_tasks = [], []
            self.teacher_data.columns = [str(c).strip() for c in self.teacher_data.columns]
            for _, row in self.teacher_data.iterrows():
                row_dict = {str(k): str(v).strip() for k, v in row.items()}
                t_name, sub, constraint, raw_classes, hours = "", "", "", "", 0
                for k, v in row_dict.items():
                    if v == "": continue
                    if '教师' in k or '老师' in k or '姓名' in k: t_name = v
                    elif '课' in k and '称' in k or '科目' in k: sub = v
                    elif '限制' in k or '连排' in k or '要求' in k: constraint = v
                    elif '班' in k: raw_classes = v
                    elif '课时' in k or '节数' in k:
                        try: hours = int(float(re.findall(r'\d+', v)[0]))
                        except: pass
                
                if not t_name and len(row_dict) >= 1: t_name = list(row_dict.values())[0]
                if not sub and len(row_dict) >= 2: sub = list(row_dict.values())[1]
                if not raw_classes and len(row_dict) >= 3: raw_classes = list(row_dict.values())[2]
                if hours <= 0 or not t_name or not raw_classes: continue

                class_nums = re.findall(r'\d+', raw_classes)
                target_classes = [f"{self.prefix}{n}{self.suffix}" for n in class_nums] 
                for c in target_classes: self.classes.add(c)

                if "Block_ZouBan" in constraint or "走班" in sub or "走班" in t_name:
                    zouban_tasks.append({'classes': target_classes, 'teacher': t_name, 'subject': sub, 'hours': hours})
                else:
                    for c in target_classes: tasks.append({'class': c, 'teacher': t_name, 'subject': sub, 'hours': hours, 'constraint': constraint})

            def sort_key(c_str):
                nums = re.findall(r'\d+', c_str)
                return int(nums[0]) if nums else 0
            self.classes = sorted(list(self.classes), key=sort_key) 
            if not self.classes: self.classes = [f'{self.prefix}1{self.suffix}'] 
            return tasks, zouban_tasks

        def is_free_for(self, c_name, teacher, subject, day, slot_list):
            for slot in slot_list:
                if self.schedules[c_name][day][slot] != "": return False 
                if (teacher, day, slot) in self.teacher_busy: return False 
            count = sum(1 for s in SLOTS if self.schedules[c_name][day][s] != "" and ("走班" in subject and "走班" in self.schedules[c_name][day][s] or self.schedules[c_name][day][s].startswith(f"{subject}<br>")))
            if count + len(slot_list) > 2: return False
            return True

        def book(self, c_name, teacher, subject, day, slot):
            content = "走班" if ("走班" in subject or "走班" in teacher) else f"{subject}<br>{teacher}"
            self.schedules[c_name][day][slot] = content
            self.teacher_busy[(teacher, day, slot)] = c_name 

        def run(self):
            self.original_tasks, self.original_zouban = self.parse_tasks()
            if not self.original_tasks and not self.original_zouban: 
                return {c: {day: {slot: "" for slot in SLOTS} for day in DAYS} for c in self.classes}, 0

            PAIRS = [(SLOTS[1], SLOTS[2]), (SLOTS[3], SLOTS[4]), (SLOTS[7], SLOTS[8])]
            best_schedule, best_score = None, 99999
            
            for attempt in range(100):
                tasks, zouban_tasks = copy.deepcopy(self.original_tasks), copy.deepcopy(self.original_zouban)
                self.schedules = {c: {day: {slot: "" for slot in SLOTS} for day in DAYS} for c in self.classes}
                self.teacher_busy = {} 

                zouban_eve_ok = True
                for z_task in zouban_tasks:
                    if z_task['hours'] < 2: continue
                    placed = False
                    days = list(self.evening_days)
                    random.shuffle(days)
                    for day in days:
                        if all(self.is_free_for(c, z_task['teacher'], z_task['subject'], day, [SLOTS[9], SLOTS[10]]) for c in z_task['classes']):
                            for c in z_task['classes']:
                                self.book(c, z_task['teacher'], z_task['subject'], day, SLOTS[9])
                                self.book(c, z_task['teacher'], z_task['subject'], day, SLOTS[10])
                            z_task['hours'] -= 2; placed = True; break
                    if not placed: zouban_eve_ok = False; break
                if not zouban_eve_ok: continue

                eve_placed_ok = True
                eve_assigned = {c: set() for c in self.classes}
                for day in self.evening_days:
                    cls_list = list(self.classes)
                    random.shuffle(cls_list)
                    for c in cls_list:
                        if self.schedules[c][day][SLOTS[9]] != "": continue 
                        c_tasks = [t for t in tasks if t['class'] == c and t['hours'] >= 2]
                        random.shuffle(c_tasks)
                        placed = False
                        for t in [t for t in c_tasks if t['subject'] not in eve_assigned[c]] + [t for t in c_tasks if t['subject'] in eve_assigned[c]]:
                            if self.is_free_for(c, t['teacher'], t['subject'], day, [SLOTS[9], SLOTS[10]]):
                                self.book(c, t['teacher'], t['subject'], day, SLOTS[9])
                                self.book(c, t['teacher'], t['subject'], day, SLOTS[10])
                                t['hours'] -= 2; eve_assigned[c].add(t['subject']); placed = True; break
                        if not placed: eve_placed_ok = False; break
                    if not eve_placed_ok: break
                if not eve_placed_ok: continue 

                for z_task in zouban_tasks:
                    while z_task['hours'] >= 2:
                        placed = False
                        all_days = DAYS[:5]
                        random.shuffle(all_days)
                        for day in all_days:
                            shuffled_pairs = PAIRS.copy()
                            random.shuffle(shuffled_pairs)
                            for s1, s2 in shuffled_pairs:
                                if all(self.is_free_for(c, z_task['teacher'], z_task['subject'], day, [s1, s2]) for c in z_task['classes']):
                                    for c in z_task['classes']:
                                        self.book(c, z_task['teacher'], z_task['subject'], day, s1)
                                        self.book(c, z_task['teacher'], z_task['subject'], day, s2)
                                    z_task['hours'] -= 2; placed = True; break
                            if placed: break
                        if not placed: break
                    
                    while z_task['hours'] > 0:
                        placed = False
                        all_times = [(d, s) for d in DAYS[:5] for s in [SLOTS[5], SLOTS[6]]]
                        random.shuffle(all_times)
                        for day, slot in all_times:
                            if all(self.is_free_for(c, z_task['teacher'], z_task['subject'], day, [slot]) for c in z_task['classes']):
                                for c in z_task['classes']: self.book(c, z_task['teacher'], z_task['subject'], day, slot)
                                z_task['hours'] -= 1; placed = True; break
                        if not placed: break

                for task in tasks:
                    while task['hours'] >= 2:
                        placed = False
                        all_days = DAYS[:5]
                        random.shuffle(all_days)
                        for day in all_days:
                            shuffled_pairs = PAIRS.copy()
                            random.shuffle(shuffled_pairs)
                            for s1, s2 in shuffled_pairs:
                                if self.is_free_for(task['class'], task['teacher'], task['subject'], day, [s1, s2]):
                                    self.book(task['class'], task['teacher'], task['subject'], day, s1)
                                    self.book(task['class'], task['teacher'], task['subject'], day, s2)
                                    task['hours'] -= 2; placed = True; break 
                            if placed: break 
                        if not placed: break 

                for task in tasks:
                    while task['hours'] > 0:
                        placed = False
                        all_days = DAYS[:5]
                        random.shuffle(all_days)
                        for day in all_days:
                            if self.is_free_for(task['class'], task['teacher'], task['subject'], day, [SLOTS[5]]):
                                self.book(task['class'], task['teacher'], task['subject'], day, SLOTS[5]); task['hours'] -= 1; placed = True; break
                        if not placed:
                            all_times = [(d, s) for d in DAYS[:5] for s in [SLOTS[1], SLOTS[2], SLOTS[3], SLOTS[4], SLOTS[7], SLOTS[8]]]
                            random.shuffle(all_times)
                            for day, slot in all_times:
                                if self.is_free_for(task['class'], task['teacher'], task['subject'], day, [slot]):
                                    self.book(task['class'], task['teacher'], task['subject'], day, slot); task['hours'] -= 1; placed = True; break
                        if not placed: break
                        
                unplaced = sum(t['hours'] for t in tasks) + sum(t['hours'] for t in zouban_tasks)
                empty_eve = sum(1 for c in self.classes for day in self.evening_days for s in [SLOTS[9], SLOTS[10]] if self.schedules[c][day][s] == "")
                score = unplaced * 10 + empty_eve
                if score < best_score:
                    best_score = score
                    best_schedule = copy.deepcopy(self.schedules)
                if score == 0: break 
                    
            self.schedules = best_schedule
            for c in self.classes:
                for day in DAYS[:5]:
                    if self.schedules[c][day][SLOTS[6]] == "": self.schedules[c][day][SLOTS[6]] = "自习"
                for day in self.evening_days:
                    for s in [SLOTS[9], SLOTS[10]]:
                        if self.schedules[c][day][s] == "": self.schedules[c][day][s] = "晚自习"

            return self.schedules, len(self.original_tasks) + len(self.original_zouban)

    def render_class_table(class_name, schedule):
        html = f"""
        <div id="table-{class_name}" style="font-family: 'SimSun', '宋体', sans-serif; max-width: 1050px; margin: auto; background-color: white; padding: 20px;">
            <h2 style="text-align: center; color: red; letter-spacing: 2px;">呼和浩特市英华学校 {class_name} 课程表</h2>
            <table border="1" style="width: 100%; text-align: center; border-collapse: collapse; border-color: #333; font-size: 16px;">
                <tr style="background-color: #f8f9fa;"><th style="padding: 12px;">时间 \\ 星期</th><th>一</th><th>二</th><th>三</th><th>四</th><th>五</th><th>六</th><th>日</th></tr>
        """
        def get_cell(day, slot):
            content = schedule[day][slot]
            bg_color = "#FFFFFF" 
            for sub, color in SUBJECT_COLORS.items():
                if sub in content: bg_color = color; break
            if "走班" in content: return f'<td style="background-color: {bg_color}; font-weight: bold; padding: 12px; box-shadow: inset 0 0 5px rgba(0,0,0,0.1);">{content}</td>'
            return f'<td style="background-color: {bg_color}; padding: 12px;">{content}</td>'

        html += "<tr><td>早自习(7:00-7:40)</td>"
        for day in DAYS[:5]: html += "<td style='padding: 12px; background-color: #F8F9FA;'>早自习</td>"
        html += "<td rowspan='12' style='width:45px; background-color: #F8F9FA; color: #555;'>考<br><br>试</td><td rowspan='12' style='width:45px; background-color: #F8F9FA; color: #555;'>休<br><br>息</td></tr>" 
        for slot in SLOTS[1:3]:
            html += f"<tr><td>{slot.replace(chr(10), '<br>')}</td>"
            for day in DAYS[:5]: html += get_cell(day, slot)
            html += "</tr>"
        html += "<tr><td colspan='6' style='letter-spacing: 25px; font-weight:bold; padding: 6px; background-color: #EFEFEF; color: #666;'>课间操</td></tr>"
        for slot in SLOTS[3:5]:
            html += f"<tr><td>{slot.replace(chr(10), '<br>')}</td>"
            for day in DAYS[:5]: html += get_cell(day, slot)
            html += "</tr>"
        html += "<tr><td colspan='6' style='letter-spacing: 25px; font-weight:bold; padding: 6px; background-color: #EFEFEF; color: #666;'>午餐</td></tr>"
        for slot in SLOTS[5:9]:
            html += f"<tr><td>{slot.replace(chr(10), '<br>')}</td>"
            for day in DAYS[:5]: html += get_cell(day, slot)
            html += "</tr>"
        html += "<tr><td colspan='8' style='letter-spacing: 25px; font-weight:bold; padding: 6px; background-color: #EFEFEF; color: #666;'>晚餐</td></tr>"
        for slot in SLOTS[9:]:
            html += f"<tr><td>{slot.replace(chr(10), '<br>')}</td>"
            for day in DAYS[:5]: html += get_cell(day, slot)
            html += "<td style='padding: 12px; background-color: #F8F9FA;'>休息</td>" + get_cell('周日', slot) + "</tr>"
        return html + "</table></div><br>"

    def export_course_to_excel(schedules, classes):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for c_name in classes:
                data = []
                for slot in SLOTS:
                    row = [slot.replace('\n', ' ')] 
                    for day in DAYS:
                        row.append(schedules[c_name][day][slot].replace("<br>", "\n"))
                    data.append(row)
                pd.DataFrame(data, columns=['时间段'] + DAYS).to_excel(writer, sheet_name=c_name, index=False)
        return output.getvalue()

    st.title("📅 全校智能排课系统")
    uploaded_course = st.file_uploader("📂 请上传【排课数据】 Excel", type=['xlsx', 'xls'], key="course_upload")

    if uploaded_course:
        all_sheets = pd.read_excel(uploaded_course, sheet_name=None)
        teacher_df = all_sheets[list(all_sheets.keys())[0]] if len(all_sheets) == 1 else [df for df in all_sheets.values() if len(df) > 5 and len(df.columns) >= 3][0]

        if st.button("🚀 开始生成教务课表", type="primary"):
            with st.spinner("正在执行算法排课，寻找完美解..."):
                scheduler = CourseScheduler(teacher_df, class_prefix, class_suffix, include_sunday_evening)
                result, total_tasks = scheduler.run()
                if total_tasks == 0: st.error("❌ 解析失败，请检查 Excel 表头。")
                else:
                    st.session_state['schedule_result'] = result
                    st.session_state['schedule_classes'] = scheduler.classes

    if st.session_state['schedule_result'] is not None:
        st.success("✅ 完美排课成功！请在下方查看结果并导出。")
        result, classes = st.session_state['schedule_result'], st.session_state['schedule_classes']
        col1, col2 = st.columns([1, 3])
        with col1:
            st.download_button(
                label="📊 导出课表到 Excel",
                data=export_course_to_excel(result, classes),
                file_name=f"排课结果_{class_prefix}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        with col2:
            components.html("""<button onclick="window.parent.print()" style="padding: 8px 15px; background-color: #FF4B4B; color: white; border: none; border-radius: 5px; cursor: pointer; font-size: 14px;">🖨️ 打印课表 / 存为彩色 PDF</button>""", height=50)

        tabs = st.tabs(classes)
        for idx, c_name in enumerate(classes):
            with tabs[idx]:
                st.markdown(render_class_table(c_name, result[c_name]), unsafe_allow_html=True)


# =====================================================================
#                          模块二：全科考场编排 (兼容多版本表头)
# =====================================================================
elif app_mode == "📝 全科考场编排":
    
    EDITOR_ORDER = [
        "语文", "数学", "物理", "历史", 
        "英语", "外语", "日语", "俄语", "西班牙语", "法语", "德语", 
        "化学", "地理", "政治", "思想政治", "生物", "生物学", "技术"
    ]
    DYNAMIC_POOL = ["化学", "地理", "政治", "思想政治", "生物", "生物学", "技术"]

    def get_exam_sort_key(ex):
        sub = str(ex.get('科目', '')).strip()
        try: return EDITOR_ORDER.index(sub)
        except ValueError: return 999  

    def format_time_display(time_str, is_html=False):
        br = '<br>' if is_html else '\n'
        time_str = str(time_str).strip()
        if ' ' in time_str: return time_str.replace(' ', br, 1)
        for kw in ['日', '号']:
            if kw in time_str:
                parts = time_str.split(kw, 1)
                return f"{parts[0]}{kw}{br}{parts[1]}"
        return time_str

    DEFAULT_EXAM_TIMES = {
        "语文": "4月22日 09:00-11:30", "数学": "4月22日 15:00-17:00",
        "物理": "4月23日 09:00-10:15", "历史": "4月23日 09:00-10:15",
        "英语": "4月23日 15:00-17:00", "外语": "4月23日 15:00-17:00",
        "日语": "4月23日 15:00-17:00", "俄语": "4月23日 15:00-17:00", "西班牙语": "4月23日 15:00-17:00",
        "化学": "4月24日 08:30-09:45", "地理": "4月24日 11:00-12:15", 
        "政治": "4月24日 14:30-15:45", "思想政治": "4月24日 14:30-15:45",
        "生物": "4月24日 17:00-18:15", "生物学": "4月24日 17:00-18:15"
    }

    with st.sidebar:
        st.header("⚙️ 考务参数设置")
        room_capacity = st.number_input("选考标准考场容量 (人数)", min_value=10, max_value=60, value=30)
        enable_balance = st.checkbox("✅ 开启选考科目考场人数均衡", value=True)
        st.markdown("---")
        st.header("🖨️ 打印排版设置")
        exam_name = st.text_input("考试名称", value="二模考试")
        slips_per_page = st.selectbox("每页打印学生数", [4, 6, 8], index=1)

    class ExamScheduler:
        def __init__(self, student_df, capacity, time_map, balance):
            self.student_data = student_df.fillna("")
            self.capacity = int(capacity)
            self.time_map = time_map
            self.balance = balance  
            self.fixed_subjects, self.dynamic_subjects = {}, {}

        # 智能表头匹配
        def find_col(self, keywords):
            for col in self.student_data.columns:
                for kw in keywords:
                    if kw in str(col): return col
            return None

        def parse_data(self):
            c_zkz = self.find_col(["准考证", "考号", "学号"])
            c_name = self.find_col(["姓名", "学生"])
            c_class = self.find_col(["班级", "行政班"])
            c_room = self.find_col(["考场"])
            c_seat = self.find_col(["座位"])

            for _, row in self.student_data.iterrows():
                student_info = {
                    '准考证号': str(row.get(c_zkz, '')), 
                    '姓名': str(row.get(c_name, '')),
                    '行政班': str(row.get(c_class, '')),
                    '原考场': str(row.get(c_room, '')).zfill(2) if row.get(c_room) else "00",
                    '原座位': str(row.get(c_seat, '')).zfill(2) if row.get(c_seat) else "00"
                }
                
                subs_for_student = ['语文', '数学']
                for col in ['语种', '科类', '选考1', '选考2', '首选科目', '再选科目']:
                    actual_col = self.find_col([col])
                    if actual_col:
                        sub = str(row.get(actual_col, '')).strip()
                        if sub and sub.lower() != 'nan': subs_for_student.append(sub)
                
                for sub in set(subs_for_student): 
                    if sub in DYNAMIC_POOL:
                        if sub not in self.dynamic_subjects: self.dynamic_subjects[sub] = []
                        self.dynamic_subjects[sub].append(student_info)
                    else:
                        if sub not in self.fixed_subjects: self.fixed_subjects[sub] = []
                        self.fixed_subjects[sub].append(student_info)

        def arrange(self):
            self.parse_data()
            exam_results, student_slips = {}, {}
            def add_to_slip(zkz, name, cls, subject, room_name, seat_name):
                if zkz not in student_slips: 
                    student_slips[zkz] = {'姓名': name, '行政班': cls, '准考证号': zkz, 'exams': []}
                time_val = self.time_map.get(subject, "时间待定")
                student_slips[zkz]['exams'].append({'科目': subject, '时间': time_val, '考场': room_name, '座位': seat_name})

            # 1. 固定科目排位
            for subject, students in self.fixed_subjects.items():
                if not students: continue 
                data_list = []
                for stu in students:
                    room_name = f"第{stu['原考场']}考场" if stu['原考场'] != "00" else "未分配"
                    add_to_slip(stu['准考证号'], stu['姓名'], stu['行政班'], subject, room_name, stu['原座位'])
                    data_list.append({'考场号': room_name, '座位号': stu['原座位'], '准考证号': stu['准考证号'], '姓名': stu['姓名'], '行政班': stu['行政班']})
                if data_list:
                    exam_results[subject] = pd.DataFrame(data_list, columns=['考场号', '座位号', '准考证号', '姓名', '行政班']).sort_values(by=['考场号', '座位号'])
                
            # 2. 选考动态排位与均衡
            for subject, students in self.dynamic_subjects.items():
                if not students: continue 
                students_sorted = sorted(students, key=lambda x: (x['原考场'], x['原座位']))
                data_list = []
                total_stu = len(students_sorted)
                num_rooms = math.ceil(total_stu / self.capacity)
                if num_rooms > 0:
                    if self.balance:
                        base_cap, rem = divmod(total_stu, num_rooms)
                        room_caps = [base_cap + (1 if i < rem else 0) for i in range(num_rooms)]
                    else:
                        room_caps = [self.capacity for _ in range(num_rooms - 1)] + [total_stu % self.capacity or self.capacity]
                else: room_caps = []

                curr_r, curr_s = 0, 0
                for stu in students_sorted:
                    if curr_s >= room_caps[curr_r]: curr_r += 1; curr_s = 0
                    rn, sn = curr_r + 1, curr_s + 1
                    room_name, seat_name = f"第{rn:02d}考场", f"{sn:02d}"
                    add_to_slip(stu['准考证号'], stu['姓名'], stu['行政班'], subject, room_name, seat_name)
                    data_list.append({'考场号': room_name, '座位号': seat_name, '准考证号': stu['准考证号'], '姓名': stu['姓名'], '行政班': stu['行政班']})
                    curr_s += 1
                if data_list:
                    exam_results[subject] = pd.DataFrame(data_list, columns=['考场号', '座位号', '准考证号', '姓名', '行政班']).sort_values(by=['考场号', '座位号'])
            
            slip_export_data = []
            for zkz, info in student_slips.items():
                row = {'准考证号': zkz, '姓名': info['姓名'], '行政班': info['行政班']}
                sorted_exams = sorted(info['exams'], key=get_exam_sort_key)
                for idx, ex in enumerate(sorted_exams):
                    row[f'科目{idx+1}'] = ex['科目']; row[f'时间{idx+1}'] = ex['时间']
                    row[f'考场{idx+1}'] = ex['考场']; row[f'座位{idx+1}'] = ex['座位']
                slip_export_data.append(row)
            return exam_results, student_slips, slip_export_data

    # ================= 准考证 Word 排版 =================
    def export_slips_to_word(slips_dict, per_page, exam_title):
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Inches(0.25), Inches(0.25), Inches(0.35), Inches(0.35)

        def set_font(run, font_name, size, bold=False, color=None):
            run.font.name = font_name; run.font.size = Pt(size); run.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if color: run.font.color.rgb = RGBColor(*color)

        student_list = list(slips_dict.values())
        num_cols, num_rows = 2, per_page // 2
        for i in range(0, len(student_list), per_page):
            grid_table = doc.add_table(rows=num_rows, cols=num_cols)
            tblPr = grid_table._tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge = OxmlElement(f'w:{b}'); edge.set(qn('w:val'), 'dashed'); edge.set(qn('w:sz'), '4'); edge.set(qn('w:color'), 'AAAAAA')
                tblBorders.append(edge)
            tblPr.append(tblBorders)
            for row in grid_table.rows:
                row.height = Inches(10.0 / num_rows); row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            chunk = student_list[i : i + per_page]
            for idx, info in enumerate(chunk):
                cell = grid_table.cell(idx // num_cols, idx % num_cols)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(8)
                tab_stops = p.paragraph_format.tab_stops; tab_stops.add_tab_stop(Inches(1.75), WD_TAB_ALIGNMENT.CENTER)
                run_s = p.add_run("🏫 英华学校\t"); set_font(run_s, '微软雅黑', 9, color=(130, 130, 130))
                run_t = p.add_run(f"{exam_title}准考证"); set_font(run_t, '黑体', 15, True)
                
                p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(8)
                run_l1 = p2.add_run("姓名: "); set_font(run_l1, '宋体', 9, False)
                run_v1 = p2.add_run(f"{info['姓名']} "); set_font(run_v1, '黑体', 10, True)
                run_l2 = p2.add_run("班级: "); set_font(run_l2, '宋体', 9, False)
                run_v2 = p2.add_run(f"{info['行政班']} "); set_font(run_v2, '黑体', 10, True)
                run_l3 = p2.add_run("考号: "); set_font(run_l3, '宋体', 9, False)
                run_v3 = p2.add_run(f"{info['准考证号']}"); set_font(run_v3, 'Arial', 10, True, color=(200, 60, 0))
                
                inner = cell.add_table(rows=1, cols=4); inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
                in_tblPr = inner._tbl.tblPr
                in_tblBorders = OxmlElement('w:tblBorders')
                for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    edge = OxmlElement(f'w:{b}'); edge.set(qn('w:val'), 'single'); edge.set(qn('w:sz'), '4'); edge.set(qn('w:color'), 'E0E0E0')
                    in_tblBorders.append(edge)
                in_tblPr.append(in_tblBorders)
                for c_idx, w in enumerate([Inches(0.6), Inches(1.4), Inches(1.0), Inches(0.5)]): inner.columns[c_idx].width = w

                hdr = inner.rows[0].cells; inner.rows[0].height = Inches(0.28); inner.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                for j, txt in enumerate(['科目', '考试时间', '考场名称', '座号']):
                    hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER; cp = hdr[j].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(cp.add_run(txt), '黑体', 9, True, color=(40, 40, 40))
                    shd = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w'))); hdr[j]._tc.get_or_add_tcPr().append(shd)

                for ex in sorted(info['exams'], key=get_exam_sort_key):
                    row_c = inner.add_row().cells; inner.rows[-1].height = Inches(0.28); inner.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    vals = [ex['科目'], format_time_display(ex['时间'], False), ex['考场'], ex['座位']]
                    for j, v in enumerate(vals):
                        row_c[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER; cp = row_c[j].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if j == 0: set_font(cp.add_run(v), '宋体', 9, color=(80, 80, 80))
                        elif j == 1: set_font(cp.add_run(v), 'Arial', 8, color=(80, 80, 80))
                        elif j == 2: set_font(cp.add_run(v), '黑体', 10, True)
                        elif j == 3: set_font(cp.add_run(v), 'Arial', 9)
            if i + per_page < len(student_list): doc.add_page_break()
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    # ================= 考场门贴 Word 排版 =================
    def export_door_signs_to_word(exam_results, exam_title):
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)
        def set_font(run, font_name, size, bold=False, color=None):
            run.font.name = font_name; run.font.size = Pt(size); run.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if color: run.font.color.rgb = RGBColor(*color)
        is_first = True
        for sub in EDITOR_ORDER:
            if sub not in exam_results or exam_results[sub].empty: continue
            df = exam_results[sub]
            for room in sorted(df['考场号'].unique()):
                if not is_first: doc.add_page_break()
                is_first = False
                p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p1.add_run(f"🏫 {exam_title} · {sub}"), '黑体', 22, True)
                p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(16)
                set_font(p2.add_run(f"【 {room} 考生名单 】"), '微软雅黑', 16, True)
                table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'; table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for c_idx, w in enumerate([Inches(1.0), Inches(2.2), Inches(1.5), Inches(1.8)]): table.columns[c_idx].width = w
                hdr = table.rows[0].cells; table.rows[0].height = Pt(32)
                for j, h in enumerate(['座位号', '准考证号', '姓名', '行政班']):
                    hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER; p = hdr[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(p.add_run(h), '黑体', 12, True)
                    shd = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w'))); hdr[j]._tc.get_or_add_tcPr().append(shd)
                for _, row in df[df['考场号'] == room].iterrows():
                    row_c = table.add_row().cells; table.rows[-1].height = Pt(26)
                    vals = [str(row['座位号']), str(row['准考证号']), str(row['姓名']), str(row['行政班'])]
                    for j, v in enumerate(vals):
                        row_c[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER; p = row_c[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_font(p.add_run(v), '黑体' if j==2 else ('Arial' if j==1 else '宋体'), 11, j==2)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    st.title("📝 全科考场编排 (教务兼容终极版)")
    st.info("💡 **系统已开启智能兼容模式**：支持高一、高二、高三所有不规范的表头格式（自动识别准考证、考号、班级等）。")
    up_exam = st.file_uploader("📂 第一步：上传【考生信息表】", type=['xlsx', 'xls', 'csv'])

    if up_exam:
        df_stu = pd.read_csv(up_exam, dtype=str) if up_exam.name.endswith('.csv') else pd.read_excel(up_exam, dtype=str)
        st.success(f"✅ 成功读取 {len(df_stu)} 名考生数据")

        st.markdown("### 📅 第二步：确认考试时间")
        all_subs = set(['语文', '数学'])
        for col in df_stu.columns:
            if any(kw in str(col) for kw in ['语种', '选考', '科类', '科目']):
                for val in df_stu[col].dropna().unique():
                    v_str = str(val).strip()
                    if v_str and v_str.lower() != 'nan': all_subs.add(v_str)
        
        time_data = [{"科目": s, "考试时间": DEFAULT_EXAM_TIMES.get(s, "时间待定")} for s in sorted(list(all_subs), key=lambda x: EDITOR_ORDER.index(x) if x in EDITOR_ORDER else 999)]
        dynamic_key = "exam_time_editor_" + hashlib.md5("".join(sorted(list(all_subs))).encode()).hexdigest()[:8]
        edited_time_df = st.data_editor(pd.DataFrame(time_data), use_container_width=True, hide_index=True, column_config={"科目": st.column_config.TextColumn("考试科目", disabled=True)}, key=dynamic_key)
        final_time_map = dict(zip(edited_time_df['科目'], edited_time_df['考试时间']))

        if st.button("🚀 第三步：生成准考证与考场数据", type="primary"):
            with st.spinner("正在启动兼容引擎，进行智能排位..."):
                sch = ExamScheduler(df_stu, room_capacity, final_time_map, enable_balance)
                res, slips, export = sch.arrange()
                st.session_state['exam_result'], st.session_state['exam_slips'], st.session_state['slip_export_data'] = res, slips, export

    if st.session_state['exam_result'] is not None:
        st.success("✅ 考务数据已全部生成完毕！")
        c1, c2, c3 = st.columns(3)
        with c1:
            out_ex = io.BytesIO()
            with pd.ExcelWriter(out_ex, engine='openpyxl') as writer:
                for sub in sorted(list(st.session_state['exam_result'].keys()), key=lambda x: EDITOR_ORDER.index(x) if x in EDITOR_ORDER else 999):
                    st.session_state['exam_result'][sub].to_excel(writer, sheet_name=f"{sub}考场", index=False)
                
                # 分卷统计表
                all_data_list = []
                for sub, df in st.session_state['exam_result'].items():
                    if not df.empty:
                        temp_df = df.copy(); temp_df['科目'] = sub; all_data_list.append(temp_df)
                if all_data_list:
                    full_df = pd.concat(all_data_list, ignore_index=True)
                    stats_df = full_df.groupby(['考场号', '科目']).size().unstack(fill_value=0)
                    ordered_cols = [c for c in EDITOR_ORDER if c in stats_df.columns]
                    stats_df = stats_df[ordered_cols]
                    stats_df['考场总卷数'] = stats_df.sum(axis=1)
                    stats_df.reset_index(inplace=True)
                    stats_df.to_excel(writer, sheet_name="分卷统计表", index=False)
                
                pd.DataFrame(st.session_state['slip_export_data']).to_excel(writer, sheet_name="个人全科汇总", index=False)
            st.download_button("📊 下载 考务总表 (Excel)", out_ex.getvalue(), f"英华_{exam_name}_考务.xlsx", use_container_width=True)
        with c2:
            word_slips = export_slips_to_word(st.session_state['exam_slips'], slips_per_page, exam_name)
            st.download_button("🎫 下载 准考证条 (Word)", word_slips, f"英华_{exam_name}_准考证.docx", use_container_width=True)
        with c3:
            word_door = export_door_signs_to_word(st.session_state['exam_result'], exam_name)
            st.download_button("🚪 下载 考场门贴 (Word)", word_door, f"英华_{exam_name}_门贴.docx", type="primary", use_container_width=True)

        tabs = st.tabs(["🎫 预览准考证", "📋 各科明细预览"])
        with tabs[0]:
            html = '<div style="display:flex; flex-wrap:wrap; justify-content:center; gap:15px; padding:10px; background:#f0f2f6;">'
            for zkz, info in list(st.session_state['exam_slips'].items())[:15]: 
                html += f"""<div style="width:340px; border:1px dashed #999; border-radius:8px; background:#fff; padding:12px; margin-bottom:10px;">
                    <div style="border-bottom:1px solid #eee; padding-bottom:6px; margin-bottom:10px; position:relative; text-align:center;">
                        <span style="font-size:10px; color:#888; position:absolute; left:0; bottom:3px;">🏫 英华学校</span><b style="font-size:17px;">{exam_name}准考证</b>
                    </div>
                    <div style="font-size:11px; margin-bottom:10px; text-align:center;">姓名: <b>{info['姓名']}</b> &nbsp; 班级: <b>{info['行政班']}</b> &nbsp; 考号: <b style="color:#cc4400; font-family:Arial;">{info['准考证号']}</b></div>
                    <table style="width:100%; border-collapse:collapse; font-size:11px; text-align:center; border:1px solid #e0e0e0;">
                        <tr style="background:#f5f5f5;"><th>科目</th><th>考试时间</th><th>考场</th><th>座号</th></tr>"""
                for ex in sorted(info['exams'], key=get_exam_sort_key):
                    html += f"<tr><td style='padding:5px; border:1px solid #e0e0e0; color:#555;'>{ex['科目']}</td><td style='line-height:1.2; padding:5px 0; border:1px solid #e0e0e0; color:#555;'>{format_time_display(ex['时间'], True)}</td><td style='font-weight:bold; border:1px solid #e0e0e0;'>{ex['考场']}</td><td style='border:1px solid #e0e0e0;'>{ex['座位']}</td></tr>"
                html += "</table></div>"
            st.markdown(html + '</div>', unsafe_allow_html=True)
            if len(st.session_state['exam_slips']) > 15: st.warning("💡 网页仅显示前 15 位考生进行预览，下载 Word 查看全校数据。")
        with tabs[1]:
            st.info("您可以切换下方标签页，快速预览各学科考场的分配情况。")
            sub_tabs = st.tabs(sorted(list(st.session_state['exam_result'].keys()), key=lambda x: EDITOR_ORDER.index(x) if x in EDITOR_ORDER else 999))
            for idx, sub in enumerate(sorted(list(st.session_state['exam_result'].keys()), key=lambda x: EDITOR_ORDER.index(x) if x in EDITOR_ORDER else 999)):
                with sub_tabs[idx]: st.dataframe(st.session_state['exam_result'][sub], use_container_width=True, hide_index=True)